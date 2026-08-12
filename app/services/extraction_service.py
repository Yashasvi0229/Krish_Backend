"""
Text extraction from PDF / DOCX / XLSX / images.

Every extractor returns an `ExtractionResult`:
    * `text` — the concatenated plain text
    * `page_count` — for PDF and DOCX only (paginated formats)
    * `ocr_applied` — whether we fell back to OCR

Why page_count matters: GNC's billing rule is "0.2 hrs per page (up to 4
pages), then 0.4 hrs per page". Without page_count, we can't bill
Document Review at all.

OCR strategy: if a PDF's text extraction returns fewer than
`OCR_MIN_CHARS_BEFORE_OCR` characters per page on average, we assume it's
scanned and rasterize + OCR. Tesseract is slow — we only invoke it when
we have to.
"""
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path

from app.config import settings
from app.core.logging import get_logger

log = get_logger(__name__)


@dataclass(frozen=True)
class ExtractionResult:
    text: str
    page_count: int | None
    ocr_applied: bool
    error: str | None = None


# ---------------------------------------------------------------------------
# Dispatcher — picks extractor by extension
# ---------------------------------------------------------------------------
def extract(data: bytes, extension: str) -> ExtractionResult:
    """Route bytes to the right extractor. Falls back to empty result on
    unsupported types — we never raise, so a single bad attachment can't
    derail an entire invoice job."""
    ext = extension.lower().lstrip(".")
    try:
        if ext == "pdf":
            return _extract_pdf(data)
        if ext in ("docx", "doc"):
            return _extract_docx(data)
        if ext in ("xlsx", "xls"):
            return _extract_xlsx(data)
        if ext in ("txt", "csv"):
            return ExtractionResult(
                text=data.decode("utf-8", errors="replace"),
                page_count=None, ocr_applied=False,
            )
        if ext in ("png", "jpg", "jpeg", "tiff", "bmp"):
            return _extract_image(data)
    except Exception as exc:  # noqa: BLE001 — extractor errors are non-fatal
        log.warning("extraction_failed", ext=ext, error=str(exc))
        return ExtractionResult(text="", page_count=None, ocr_applied=False, error=str(exc))

    return ExtractionResult(
        text="", page_count=None, ocr_applied=False,
        error=f"Unsupported extension: {ext}"
    )


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def _extract_pdf(data: bytes) -> ExtractionResult:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages = reader.pages
    page_texts: list[str] = []
    for p in pages:
        try:
            page_texts.append(p.extract_text() or "")
        except Exception as exc:  # noqa: BLE001
            log.debug("pdf_page_extract_failed", error=str(exc))
            page_texts.append("")

    text = "\n\n".join(page_texts).strip()
    page_count = len(pages)

    # Decide if we need OCR: too little text for the page count?
    threshold = settings.ocr_min_chars_before_ocr * max(page_count, 1)
    if len(text) < threshold:
        ocr_text = _ocr_pdf(data)
        if ocr_text and len(ocr_text) > len(text):
            return ExtractionResult(text=ocr_text, page_count=page_count, ocr_applied=True)

    return ExtractionResult(text=text, page_count=page_count, ocr_applied=False)


def _ocr_pdf(data: bytes) -> str:
    """Rasterize PDF to images, run Tesseract on each page.

    Returns "" if pdf2image / tesseract aren't installed on the host —
    which is fine on Render free tier without the binaries.
    """
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError:
        log.info("ocr_libs_missing_skipping")
        return ""

    try:
        images = convert_from_bytes(data, dpi=200)
    except Exception as exc:  # noqa: BLE001
        log.warning("pdf2image_failed", error=str(exc))
        return ""

    if settings.tesseract_cmd and Path(settings.tesseract_cmd).exists():
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd

    texts: list[str] = []
    for i, img in enumerate(images):
        try:
            texts.append(pytesseract.image_to_string(img, lang=settings.tesseract_lang))
        except Exception as exc:  # noqa: BLE001
            log.warning("tesseract_failed_page", page=i, error=str(exc))
            texts.append("")
    return "\n\n".join(texts).strip()


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def _extract_docx(data: bytes) -> ExtractionResult:
    from docx import Document

    doc = Document(BytesIO(data))
    # Body paragraphs
    parts = [p.text for p in doc.paragraphs if p.text]
    # Tables (invoices/reports often live in tables)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    # DOCX doesn't have a real page count in the XML — Word computes it at
    # render time. `sectPr` gives us section count as a rough proxy; if we
    # can't determine, return None and let downstream code use paragraph
    # count as a heuristic if needed.
    text = "\n".join(parts).strip()
    page_count = _estimate_docx_pages(doc)
    return ExtractionResult(text=text, page_count=page_count, ocr_applied=False)


def _estimate_docx_pages(doc) -> int:
    """Heuristic: ~500 words per page for typical letters. Not exact,
    but stable enough for billing when paired with a manual override."""
    words = 0
    for p in doc.paragraphs:
        words += len(p.text.split())
    return max(1, (words + 499) // 500)


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------
def _extract_xlsx(data: bytes) -> ExtractionResult:
    from openpyxl import load_workbook

    wb = load_workbook(BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"### Sheet: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append("\t".join(cells))
    text = "\n".join(parts).strip()
    # "pages" doesn't really apply to spreadsheets — we return sheet count.
    return ExtractionResult(text=text, page_count=len(wb.worksheets), ocr_applied=False)


# ---------------------------------------------------------------------------
# Images (photos of documents — direct OCR)
# ---------------------------------------------------------------------------
def _extract_image(data: bytes) -> ExtractionResult:
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        return ExtractionResult(text="", page_count=None, ocr_applied=False,
                                error="OCR libs not installed")
    img = Image.open(BytesIO(data))
    if settings.tesseract_cmd and Path(settings.tesseract_cmd).exists():
        pytesseract.pytesseract.tesseract_cmd = settings.tesseract_cmd
    text = pytesseract.image_to_string(img, lang=settings.tesseract_lang).strip()
    return ExtractionResult(text=text, page_count=1, ocr_applied=True)
